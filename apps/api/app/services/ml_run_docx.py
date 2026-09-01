"""Admin-only DOCX rendering from the persisted canonical technical report."""

from __future__ import annotations

from io import BytesIO
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(90, 101, 115)
HEADER_FILL = "E8EEF5"


def _set_font(run, *, size: float = 11, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("DOCX table widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _text(value: Any) -> str:
    if value is None or value == "":
        return "—"
    if isinstance(value, float):
        return f"{value:.6g}"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (list, tuple)):
        return ", ".join(_text(item) for item in value) if value else "None"
    return str(value)


def _metric_text(metrics: dict[str, Any] | None) -> str:
    if not metrics:
        return "—"
    return "; ".join(f"{name}: {_text(value)}" for name, value in sorted(metrics.items()) if not isinstance(value, dict))


def _add_table(document: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths: list[int]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        _shade(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            _set_font(run, size=9, bold=True, color=INK)
    _repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = _text(value)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _set_font(run, size=9)
    _set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _add_key_values(document: Document, rows: list[tuple[str, Any]]) -> None:
    _add_table(document, ["Field", "Value"], rows, [2700, 6660])


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, INK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("DCLab  |  ML Run Technical Report")
    _set_font(run, size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Admin technical evidence")
    _set_font(run, size=8, color=MUTED)


def render_ml_run_report_docx(report: dict[str, Any]) -> bytes:
    """Render a persisted ``RunTechnicalReport`` payload; no frontend state."""
    document = Document()
    _configure_document(document)
    run_info = dict(report.get("run") or {})
    dataset = dict(report.get("dataset") or {})

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("DCLab ML Run Report")
    _set_font(title_run, size=24, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(f"Run {run_info.get('run_id', '—')}  |  {dataset.get('name', 'Dataset')}")
    _set_font(subtitle_run, size=10.5, color=MUTED)

    _add_heading(document, "Run summary")
    _add_key_values(
        document,
        [
            ("Status", run_info.get("status")),
            ("Experiment", run_info.get("experiment_id")),
            ("Started", run_info.get("started_at")),
            ("Completed", run_info.get("completed_at")),
            ("Duration (seconds)", run_info.get("duration_seconds")),
            ("Last successful stage", run_info.get("last_successful_stage")),
            ("Failed stage", run_info.get("failed_stage")),
            ("Failure reason", run_info.get("failure_reason")),
        ],
    )

    profile = dict(report.get("raw_profile") or {})
    _add_heading(document, "Dataset analysis")
    _add_key_values(
        document,
        [
            ("Dataset", dataset.get("name")),
            ("Category", dataset.get("category")),
            ("Records uploaded", dataset.get("record_count")),
            ("Profiled rows", profile.get("row_count")),
            ("Columns", profile.get("column_count")),
            ("Missing cells", profile.get("missing_count")),
            ("Duplicate rows", profile.get("duplicate_rows", profile.get("duplicate_count"))),
            ("Constant columns", profile.get("constant_columns")),
            ("High-cardinality columns", profile.get("high_cardinality_columns")),
        ],
    )

    quality = dict(report.get("data_quality") or {})
    _add_heading(document, "Data-quality findings")
    _add_key_values(
        document,
        [
            ("Rows assessed", quality.get("row_count", profile.get("row_count"))),
            ("Issue count", quality.get("issue_count")),
            ("Findings", quality.get("issues")),
            ("Audit log", quality.get("log")),
        ],
    )

    target = dict(report.get("target_decision") or {})
    task = dict(report.get("task") or {})
    _add_heading(document, "Target and task decision")
    _add_key_values(
        document,
        [
            ("Target", target.get("column", task.get("target"))),
            ("Task", target.get("task_type", task.get("task_type"))),
            ("Metric", target.get("evaluation_metric", task.get("evaluation_metric"))),
            ("Source", target.get("source")),
            ("Confidence", target.get("confidence")),
            ("Reason", target.get("reason")),
            ("Validator", target.get("validator_verdict")),
        ],
    )

    split = dict(report.get("split") or {})
    _add_heading(document, "Locked train/test split")
    _add_key_values(
        document,
        [
            ("Strategy", split.get("strategy")),
            ("Training rows", split.get("n_train")),
            ("Test rows", split.get("n_test")),
            ("Test size", split.get("test_size")),
            ("Random state", split.get("random_state")),
            ("Stratified", split.get("stratify")),
            ("Provenance", split.get("provenance_column")),
            ("Partitions disjoint", split.get("provenance_disjoint")),
        ],
    )

    cleaning = dict(report.get("cleaning") or {})
    _add_heading(document, "Cleaning decisions")
    transformations = list(cleaning.get("transformations") or [])
    _add_table(
        document,
        ["Action", "Columns / result", "Scope"],
        [
            (
                item.get("step"),
                item.get("columns") or f"{item.get('rows_removed', item.get('cells_cleared', 0))} affected",
                cleaning.get("scope", "structural"),
            )
            for item in transformations
        ]
        or [("None required", "No structural changes", cleaning.get("scope"))],
        [2600, 4300, 2460],
    )
    plan = dict(cleaning.get("missing_value_plan") or {})
    decisions = list(plan.get("column_decisions") or [])
    _add_table(
        document,
        ["Column", "Missing", "Rate", "Decision"],
        [
            (item.get("column"), item.get("missing_count"), item.get("missing_fraction"), item.get("action"))
            for item in decisions
        ]
        or [("—", 0, 0, "No missing-value action required")],
        [2600, 1400, 1400, 3960],
    )

    roles = dict(report.get("column_roles") or {})
    _add_heading(document, "Column roles")
    _add_key_values(document, [(name.replace("_", " ").title(), value) for name, value in roles.items()])

    features = dict(report.get("feature_engineering") or {})
    _add_heading(document, "Feature engineering")
    _add_key_values(
        document,
        [
            ("Original features", features.get("original_features")),
            ("Generated features", features.get("generated_features")),
            ("Transformed features", features.get("transformed_features")),
            ("Removed features", features.get("removed_features")),
            ("Actions", features.get("feature_engineering_actions")),
        ],
    )

    preprocessing = dict(report.get("preprocessing") or {})
    _add_heading(document, "Preprocessing")
    _add_key_values(document, [(name.replace("_", " ").title(), value) for name, value in preprocessing.items()])

    candidates = list(report.get("candidate_models") or [])
    _add_heading(document, "Candidate models")
    selected_id = (report.get("selection") or {}).get("selected_candidate_id")
    for index, candidate in enumerate(candidates, start=1):
        _add_heading(document, f"{index}. {candidate.get('model_family', 'Candidate')}", level=2)
        _add_key_values(
            document,
            [
                ("Candidate ID", candidate.get("candidate_id")),
                ("Status", candidate.get("status")),
                ("Selected", candidate.get("candidate_id") == selected_id),
                ("Hyperparameters", candidate.get("hyperparameters")),
                ("Feature set", candidate.get("feature_set", candidate.get("features"))),
                ("Preprocessing", candidate.get("preprocessing_config")),
                ("CV mean", _metric_text(candidate.get("cv_mean"))),
                ("CV standard deviation", _metric_text(candidate.get("cv_std"))),
                ("Fit duration (ms)", candidate.get("fit_duration_ms")),
                ("Failure reason", candidate.get("failure_reason")),
            ],
        )
        fold_rows = [
            (fold_index, _metric_text(metrics))
            for fold_index, metrics in enumerate(candidate.get("fold_metrics") or [], start=1)
        ]
        if fold_rows:
            _add_table(document, ["Fold", "Validation metrics"], fold_rows, [1200, 8160])

    selection = dict(report.get("selection") or {})
    _add_heading(document, "Winner selection")
    _add_key_values(
        document,
        [
            ("Selected candidate", selection.get("selected_candidate_id")),
            ("Selection metric", selection.get("selection_metric")),
            ("Evidence source", selection.get("selection_source")),
            ("Locked before holdout", selection.get("locked")),
            ("Locked at", selection.get("locked_at")),
        ],
    )

    final_model = dict(report.get("final_model") or {})
    final_fit = dict(report.get("final_fit") or {})
    final_test = dict(report.get("final_test_evaluation") or {})
    _add_heading(document, "Final winner and holdout evaluation")
    _add_key_values(
        document,
        [
            ("Model family", final_model.get("model_family")),
            ("Candidate ID", final_test.get("candidate_id")),
            ("Final fit started", final_fit.get("started_at")),
            ("Final fit completed", final_fit.get("ended_at")),
            ("Final fit duration (ms)", final_fit.get("duration_ms")),
            ("Holdout evaluations", final_test.get("evaluation_count")),
            ("Test evaluation started", final_test.get("started_at")),
            ("Test evaluation completed", final_test.get("ended_at")),
            ("Evaluation duration (ms)", final_test.get("duration_ms")),
            ("Final test metrics", _metric_text(final_test.get("metrics"))),
        ],
    )

    predictions = dict(report.get("predictions_summary") or {})
    _add_heading(document, "Predictions")
    _add_key_values(document, [("Prediction count", predictions.get("count")), ("Artifact", predictions.get("artifact"))])

    timings = list(report.get("stage_timings") or [])
    _add_heading(document, "Stage timings")
    _add_table(
        document,
        ["Stage", "Duration ms", "Rows in", "Rows out", "Status"],
        [
            (item.get("stage"), item.get("duration_ms"), item.get("rows_in"), item.get("rows_out"), item.get("status"))
            for item in timings
        ],
        [2600, 1800, 1400, 1400, 2160],
    )

    verification = dict(report.get("deterministic_verification") or {})
    _add_heading(document, "Pipeline verification")
    _add_key_values(
        document,
        [
            ("Overall status", verification.get("overall_status")),
            ("Summary", verification.get("summary")),
            ("Checks", len(verification.get("checks") or [])),
            ("Warnings", len(verification.get("warnings") or [])),
            ("Failures", len(verification.get("failures") or [])),
            ("Missing evidence", len(verification.get("missing_evidence") or [])),
        ],
    )

    _add_heading(document, "Stage verification table")
    _add_table(
        document,
        ["Stage", "Status", "Checks"],
        [
            (item.get("stage"), item.get("status"), item.get("check_ids"))
            for item in (verification.get("stages") or [])
        ]
        or [("—", "NOT_VERIFIABLE", "No deterministic stage evidence")],
        [2800, 2200, 4360],
    )

    _add_heading(document, "Verification warnings")
    _add_table(
        document,
        ["Check", "Stage", "Message"],
        [
            (item.get("check_id"), item.get("stage"), item.get("message"))
            for item in [
                *(verification.get("warnings") or []),
                *(verification.get("missing_evidence") or []),
            ]
        ]
        or [("—", "—", "No deterministic verification warnings")],
        [2500, 2200, 4660],
    )

    _add_heading(document, "Verification failures")
    _add_table(
        document,
        ["Check", "Stage", "Message"],
        [
            (item.get("check_id"), item.get("stage"), item.get("message"))
            for item in (verification.get("failures") or [])
        ]
        or [("—", "—", "No deterministic verification failures")],
        [2500, 2200, 4660],
    )

    audit = list(report.get("decision_records") or [])
    _add_heading(document, "Decision and LLM audit")
    _add_table(
        document,
        ["Column", "Source", "Rule", "Final", "Validator"],
        [
            (
                item.get("column"),
                item.get("source"),
                item.get("rule_decision"),
                item.get("final_decision"),
                item.get("validator_verdict"),
            )
            for item in audit
        ]
        or [("—", "rule", "—", "—", "No semantic decisions recorded")],
        [1900, 1200, 1900, 1900, 2460],
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()
