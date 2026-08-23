from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from pathlib import Path

out = Path("Decision_AI_Agent_Coding_Context.docx")

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_heading("Decision.ai — Complete AI Coding Agent Context", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Master product, ML architecture, engineering, implementation, and validation context")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("How to use this document", level=1)
doc.add_paragraph(
    "This document is the source-of-truth context for an AI coding agent that will implement Decision.ai. "
    "The agent must read and understand the entire document before making architectural or implementation decisions. "
    "Do not reduce the project to an AutoML system or a collection of models. The objective is a multi-layer "
    "business decision-intelligence engine that explores feature and model space, creates reliable intelligence "
    "states, connects prediction to recommendation and outcome modeling, supports scenario/counterfactual analysis, "
    "and learns from real-world decisions and outcomes."
)

doc.add_heading("1. Product identity and core thesis", level=1)
doc.add_paragraph(
    "Decision.ai is an AI/ML decision-intelligence platform for businesses. It is designed to turn business data "
    "into predictive intelligence, recommendations, simulations, decisions, actions, and measurable feedback."
)
doc.add_paragraph("Core loop:")
doc.add_paragraph(
    "OBSERVE → UNDERSTAND → PREDICT → RECOMMEND → SIMULATE → DECIDE → ACT → MEASURE → LEARN → OBSERVE"
)
doc.add_paragraph(
    "The disruption thesis is not that 100 models are inherently intelligent. The thesis is that a well-engineered "
    "system can automatically explore a much larger hypothesis space than a normal manual data-science workflow, "
    "find strong and complementary models/features, compose business-specific intelligence layers, and reduce "
    "time-to-useful decision intelligence from weeks to hours or days."
)

doc.add_heading("2. Non-negotiable correction to the '100 models' concept", level=1)
doc.add_paragraph(
    "Never implement '100 random models = intelligence'. A layer may evaluate hundreds of candidate configurations, "
    "but it must retain only models supported by evidence. Selection should optimize predictive quality, calibration, "
    "temporal robustness, subgroup robustness, stability, diversity, and business relevance."
)
doc.add_paragraph(
    "A typical layer may evaluate 100–500+ candidates and retain approximately 20–50 strong, diverse models when "
    "that improves the ensemble. This is not a fixed number. If 8 models are sufficient, retain 8. If 35 are useful, "
    "retain 35. The system must justify model count through validation evidence."
)

doc.add_heading("3. Core hierarchy", level=1)
for x in [
    "Model → Intelligence Layer → Intelligence Domain → Intelligence State → Cross-Layer Decision Graph → Simulation → Decision → Action → Outcome → Feedback",
    "Model: atomic predictive component.",
    "Intelligence Layer: one meaningful business question, such as purchase probability.",
    "Intelligence Domain: related layers, such as Prediction, Recommendation, Outcome, or Simulation.",
    "Intelligence State: consolidated, versioned output of a layer, including predictions, confidence, model lineage, data quality, and evidence.",
    "Decision Graph: semantic connections between intelligence states.",
    "Simulation: evaluates possible actions and expected consequences.",
    "Decision: selects an action according to objectives, constraints, uncertainty, and policy.",
    "Feedback: records what actually happened and feeds evaluation and future learning."
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("4. Prediction domain", level=1)
doc.add_paragraph("Initial example layers:")
for x in [
    "purchase_probability",
    "churn_probability",
    "upsell_probability",
    "cross_sell_probability",
    "discount_sensitivity",
    "price_sensitivity",
    "engagement_probability",
    "email_response",
    "campaign_response",
    "next_purchase_time"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("5. Recommendation domain", level=1)
for x in [
    "send_email", "call_customer", "offer_discount", "change_price", "change_product",
    "retarget", "cross_sell", "upsell", "do_nothing", "change_channel", "change_message"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_paragraph(
    "Recommendation layers must predict the consequences and value of actions. For example, an email action can "
    "have models for probability of success, conversion, expected revenue, expected margin, expected cost, "
    "engagement, churn impact, and customer value."
)

doc.add_heading("6. Outcome domain", level=1)
for x in [
    "purchase outcome", "revenue outcome", "margin outcome", "churn outcome",
    "engagement outcome", "customer lifetime value", "customer reaction", "campaign performance"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("7. Feature intelligence", level=1)
doc.add_paragraph(
    "Feature intelligence is one of the central ideas. The system must search meaningful combinations of information "
    "rather than merely selecting one algorithm. Feature groups can include customer/demographic, behavioral, "
    "transactional, marketing, product, temporal, economic, sequence, contextual, and historical-intervention data."
)
doc.add_paragraph("Feature exploration should include:")
for x in [
    "feature-group combinations",
    "recency/frequency/monetary features",
    "lags",
    "rolling windows",
    "trend features",
    "ratios",
    "interactions",
    "aggregations",
    "time-aware transformations",
    "domain-specific transformations"
]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph("The engine must actively prevent:")
for x in [
    "target leakage",
    "future-information leakage",
    "post-action features being used to predict pre-action decisions",
    "duplicate or near-duplicate features",
    "unstable features",
    "features unavailable at inference time",
    "features that violate business or privacy constraints"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("8. Candidate model factory", level=1)
for x in [
    "Logistic Regression",
    "Elastic Net / regularized linear models",
    "Random Forest",
    "Extra Trees",
    "Gradient Boosting",
    "LightGBM",
    "XGBoost",
    "CatBoost when categorical structure warrants it",
    "calibrated variants",
    "stacking",
    "weighted blending"
]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph(
    "The model factory should generate reproducible candidate specifications consisting of model family, "
    "feature-set version, preprocessing version, hyperparameters, random seed, dataset version, and experiment ID."
)

doc.add_heading("9. Evaluation methodology", level=1)
doc.add_paragraph(
    "The system must never select models using training performance alone. Validation must respect the business "
    "prediction problem and time direction. For temporal business problems, prefer temporal train/validation/test "
    "splits and, where appropriate, rolling or expanding-window evaluation."
)
doc.add_paragraph("Binary classification metrics:")
for x in ["PR-AUC", "ROC-AUC", "precision", "recall", "F1", "Brier score", "calibration error"]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph("Also evaluate:")
for x in [
    "temporal robustness",
    "subgroup performance",
    "prediction stability",
    "calibration",
    "data quality",
    "drift sensitivity",
    "prediction diversity",
    "business relevance"
]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph(
    "Do not optimize blindly for accuracy. For imbalanced business targets, PR-AUC and calibration may matter more "
    "than raw accuracy. The layer configuration must define the primary metric and any minimum guardrails."
)

doc.add_heading("10. Diversity-aware model selection", level=1)
doc.add_paragraph(
    "The ensemble should not contain 30 models that all make essentially identical predictions. After performance "
    "screening, compare out-of-fold predictions and errors. Use prediction correlation, error correlation, algorithm "
    "family diversity, feature-space diversity, temporal robustness, and calibration to select complementary models."
)
doc.add_paragraph(
    "A strong selection pipeline is: candidate generation → leakage checks → training → out-of-fold predictions → "
    "evaluation → quality filtering → diversity analysis → multi-objective selection → ensemble."
)

doc.add_heading("11. Ensemble design", level=1)
for x in [
    "weighted probability blending",
    "stacking using out-of-fold predictions",
    "calibrated ensemble outputs",
    "model weights based on validation evidence",
    "diversity-aware weighting",
    "ensemble stability checks"
]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph(
    "Never average models simply because there are many of them. The ensemble must be compared against the best "
    "single model. If the ensemble does not improve the relevant objective or robustness, do not force it."
)

doc.add_heading("12. Intelligence State", level=1)
doc.add_paragraph("Every layer should produce a structured Intelligence State containing at least:")
for x in [
    "entity ID",
    "layer name and version",
    "prediction/output",
    "confidence or uncertainty where available",
    "models evaluated",
    "models selected",
    "ensemble method",
    "model versions",
    "feature-set versions",
    "data version",
    "primary validation metrics",
    "model agreement/disagreement",
    "data-quality indicators",
    "drift indicators",
    "important signals/explanations",
    "timestamp",
    "lineage"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("13. Cross-layer reasoning", level=1)
doc.add_paragraph(
    "Layer outputs must remain semantically meaningful. Do not immediately concatenate every layer into one opaque "
    "model. The platform should preserve the meaning and provenance of each state and allow a decision graph to "
    "combine them."
)
doc.add_paragraph("Example:")
doc.add_paragraph(
    "Purchase probability = 0.78; churn probability = 0.08; upsell probability = 0.61; "
    "discount sensitivity = 0.82; email response = 0.73."
)
doc.add_paragraph(
    "These states can inform a recommendation such as an upsell email rather than automatically applying a discount."
)

doc.add_heading("14. Recommendation intelligence", level=1)
doc.add_paragraph(
    "Recommendation is not simply classification of 'best action'. Each candidate action can be evaluated across "
    "multiple objectives."
)
for x in [
    "probability of success",
    "expected conversion",
    "expected revenue",
    "expected margin",
    "expected cost",
    "expected churn reduction/increase",
    "expected engagement",
    "expected customer lifetime value",
    "risk",
    "confidence/uncertainty"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("15. Scenario simulation", level=1)
doc.add_paragraph("The system should represent scenarios such as:")
for x in [
    "call customer",
    "send email",
    "offer 10% discount",
    "offer 20% discount",
    "retarget",
    "change channel",
    "change message",
    "upsell",
    "cross-sell",
    "do nothing"
]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph(
    "For each scenario, the engine predicts relevant outcomes, propagates uncertainty where possible, calculates "
    "expected business value, applies constraints, and compares alternatives."
)

doc.add_heading("16. Critical causal/counterfactual requirement", level=1)
doc.add_paragraph(
    "Ordinary predictive ML estimates associations such as P(Y | X). It does not automatically estimate what would "
    "happen if the company intervened with action A. A simulation asking 'what if we give a discount?' is a "
    "counterfactual/causal problem. The system must therefore distinguish predictive models from treatment-effect "
    "models."
)
for x in [
    "randomized experiments",
    "treatment/control data",
    "propensity scores",
    "uplift modeling",
    "treatment-effect estimation",
    "causal forests",
    "doubly robust estimators",
    "policy evaluation"
]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph(
    "Never claim causal simulation accuracy from ordinary observational prediction alone."
)

doc.add_heading("17. Decision engine", level=1)
doc.add_paragraph(
    "The decision engine converts predicted consequences into a recommended action according to an explicit objective "
    "function and constraints."
)
doc.add_paragraph("Example objective: maximize expected margin.")
doc.add_paragraph("Example constraints: discount <= 15%, contact frequency <= 2/7 days, customer eligible = true, budget <= configured limit.")
doc.add_paragraph(
    "The final decision must include the recommended action, expected value, confidence/uncertainty, alternatives, "
    "and machine-readable reasons/lineage."
)

doc.add_heading("18. Feedback loop", level=1)
doc.add_paragraph(
    "Every decision should generate a decision trace so the platform can compare prediction with reality."
)
for x in [
    "decision_id",
    "entity_id",
    "state snapshot",
    "recommendation",
    "action actually executed",
    "timestamp",
    "treatment/experiment assignment",
    "model versions",
    "feature versions",
    "observed outcome",
    "financial outcome",
    "human override",
    "override reason"
]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph(
    "The feedback dataset becomes a major asset: decision → action → outcome → model evaluation → improved policy."
)

doc.add_heading("19. Product positioning for B2B", level=1)
doc.add_paragraph(
    "Do not initially position Decision.ai as a replacement for internal data scientists. Most serious companies "
    "already have ML/data teams. The initial value proposition should be an engineering accelerator and decision "
    "intelligence layer that reduces repetitive experimentation and time-to-decision."
)
doc.add_paragraph(
    "Target buyers/users can include Heads of Data Science, AI/ML leaders, Data/Analytics leaders, CDO organizations, "
    "and revenue/marketing analytics teams."
)
doc.add_paragraph(
    "The strongest proof is not a slide claiming '100 models'. It is a controlled benchmark showing that Decision.ai "
    "can explore more hypotheses, produce equal or better predictive quality, and reduce engineering time and/or "
    "compute cost for a real business problem."
)

doc.add_heading("20. First vertical slice — mandatory starting point", level=1)
doc.add_paragraph(
    "Do not build all domains simultaneously. The first complete implementation is one real prediction layer: "
    "purchase_probability."
)
for x in [
    "ingest a real dataset",
    "profile schema and data quality",
    "define target and prediction horizon",
    "perform temporal split",
    "identify feature groups",
    "generate feature combinations",
    "generate candidate model configurations",
    "train candidates",
    "generate out-of-fold predictions",
    "evaluate candidates",
    "detect leakage",
    "measure diversity",
    "select strong/diverse candidates",
    "build ensemble",
    "calibrate output",
    "compare ensemble to best single model",
    "persist experiment and lineage",
    "produce Intelligence State",
    "expose result through API"
]:
    doc.add_paragraph(x, style="List Number")

doc.add_heading("21. First benchmark", level=1)
doc.add_paragraph(
    "Compare a manually designed strong baseline against Decision.ai. Track time, candidate count, compute, feature "
    "count, PR-AUC, ROC-AUC, precision, recall, F1, calibration, robustness, and business outcome metrics."
)
doc.add_paragraph(
    "The key product metric is time-to-useful-intelligence, not number of models."
)

doc.add_heading("22. Recommended software architecture", level=1)
doc.add_paragraph(
    "Start as a modular monolith. Do not prematurely create dozens of microservices. Keep clean module boundaries so "
    "services can be extracted later when scaling requires it."
)
for x in [
    "API layer",
    "configuration",
    "data ingestion",
    "dataset profiling",
    "feature engineering",
    "feature registry",
    "model factory",
    "training orchestration",
    "evaluation",
    "diversity selection",
    "ensemble",
    "layer registry",
    "domain registry",
    "intelligence-state store",
    "decision graph",
    "recommendation engine",
    "outcome engine",
    "causal engine",
    "simulation engine",
    "decision engine",
    "feedback/event ingestion",
    "monitoring",
    "model registry"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("23. Suggested technology direction", level=1)
doc.add_paragraph(
    "Python is the primary ML/backend language. FastAPI is suitable for APIs. PostgreSQL is the initial metadata and "
    "transactional store. Object storage should hold datasets and model artifacts. MLflow can provide experiment/model "
    "tracking. For distributed workloads, the orchestration layer can later use queues/workers and cloud-native "
    "compute. The exact cloud should remain configurable."
)

doc.add_heading("24. Core data objects", level=1)
for x in [
    "Tenant",
    "Dataset",
    "DatasetVersion",
    "FeatureDefinition",
    "FeatureSet",
    "FeatureSetVersion",
    "Domain",
    "Layer",
    "LayerVersion",
    "ModelCandidate",
    "ModelRun",
    "ModelArtifact",
    "Ensemble",
    "IntelligenceState",
    "DecisionGraph",
    "Scenario",
    "ScenarioResult",
    "Decision",
    "Action",
    "Experiment",
    "Treatment",
    "Outcome",
    "FeedbackEvent",
    "MonitoringMetric",
    "DriftEvent"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("25. Lineage requirement", level=1)
doc.add_paragraph(
    "Every final decision must be traceable backward. A user must be able to answer: Why was this recommendation "
    "made? Which intelligence states were used? Which models produced them? Which feature sets and dataset version "
    "were used? What validation evidence supported the models? What action happened? What was the outcome?"
)
doc.add_paragraph(
    "Decision → recommendation → intelligence states → ensemble → selected models → feature sets → features → "
    "dataset version must be traceable."
)

doc.add_heading("26. MLOps and production requirements", level=1)
for x in [
    "version every dataset, feature set, layer, model, ensemble, and decision policy",
    "reproducible training",
    "experiment tracking",
    "model registry",
    "data validation",
    "schema validation",
    "drift detection",
    "performance monitoring",
    "calibration monitoring",
    "latency monitoring",
    "cost monitoring",
    "audit logs",
    "rollback capability",
    "model approval workflow",
    "tenant isolation",
    "RBAC",
    "secrets management",
    "encryption",
    "observability"
]:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("27. Engineering rules for the coding agent", level=1)
rules = [
    "Do not invent business meaning for a feature without documenting the assumption.",
    "Do not use future information in training or inference.",
    "Do not leak target-derived information into features.",
    "Do not optimize only for accuracy.",
    "Do not force an ensemble when a single model is better.",
    "Do not claim causal effects from ordinary predictive models.",
    "Do not create 100 random models merely to satisfy a number.",
    "Prefer strong, diverse, reproducible candidates.",
    "Write tests for every non-trivial ML component.",
    "Keep experiment configuration separate from implementation.",
    "Make every experiment reproducible with a seed and dataset/model versions.",
    "Persist model lineage.",
    "Use out-of-fold predictions for stacking/ensemble evaluation.",
    "Use temporal validation when the business problem is temporal.",
    "Treat data quality as a first-class input.",
    "Do not hide uncertainty.",
    "Keep modules replaceable so algorithms can evolve.",
    "Build the first vertical slice completely before expanding breadth.",
    "Every new capability must have a measurable validation criterion."
]
for x in rules:
    doc.add_paragraph(x, style="List Bullet")

doc.add_heading("28. Development roadmap", level=1)
roadmap = [
    ("Phase 0", "Foundation: repository, configuration, API, database, experiment tracking."),
    ("Phase 1", "Dataset intelligence and prediction factory."),
    ("Phase 2", "Feature-space and model-space exploration."),
    ("Phase 3", "Evaluation, calibration, diversity selection, ensemble."),
    ("Phase 4", "Purchase probability layer and additional prediction layers."),
    ("Phase 5", "Cross-layer intelligence state and decision graph."),
    ("Phase 6", "Recommendation factory."),
    ("Phase 7", "Outcome factory."),
    ("Phase 8", "Causal/counterfactual engine."),
    ("Phase 9", "Simulation engine."),
    ("Phase 10", "Decision and policy engine."),
    ("Phase 11", "Action execution and feedback."),
    ("Phase 12", "Monitoring, retraining, governance, enterprise deployment.")
]
for phase, text in roadmap:
    doc.add_paragraph(f"{phase}: {text}", style="List Bullet")

doc.add_heading("29. Definition of success for the first MVP", level=1)
doc.add_paragraph(
    "A real business dataset can be supplied to Decision.ai with a target definition and prediction horizon. "
    "The platform automatically discovers candidate feature sets, trains and evaluates many model configurations, "
    "selects strong/diverse models, builds and validates an ensemble, creates a versioned Intelligence State, and "
    "reports the complete experiment lineage. The result must be reproducible and demonstrably competitive with a "
    "strong manually engineered baseline."
)

doc.add_heading("30. Long-term platform vision", level=1)
doc.add_paragraph(
    "The finished system should behave like a machine-learning decision operating system for business: it continuously "
    "understands data, generates predictive intelligence, composes multiple intelligence layers, evaluates candidate "
    "actions, models expected consequences, recommends decisions, records execution, measures real outcomes, and "
    "improves future decisions using the evidence generated by previous decisions."
)
doc.add_paragraph(
    "The long-term moat is therefore not the use of any individual algorithm. It is the integrated system of feature "
    "intelligence, model orchestration, diversity-aware ensemble learning, semantic intelligence layers, cross-layer "
    "reasoning, causal evidence, simulation, decision policies, outcome measurement, lineage, and accumulated "
    "decision-to-outcome knowledge."
)

doc.add_heading("31. Immediate instruction to the coding agent", level=1)
doc.add_paragraph(
    "Start by inspecting the existing repository and this context. Do not rewrite the architecture from scratch. "
    "First implement and test the complete purchase_probability vertical slice. Build the dataset profiling and "
    "feature intelligence components before adding broad prediction/recommendation domains. At each step, explain "
    "the engineering decision, add tests, preserve lineage, and validate the result against a baseline. Only after "
    "the first layer is working end-to-end should the agent generalize it into a reusable Layer/Domain framework."
)

doc.add_heading("32. Final product statement", level=1)
doc.add_paragraph(
    "Decision.ai builds AI that does not stop at predicting what may happen. It builds structured intelligence from "
    "many validated perspectives, connects predictions to possible actions, estimates consequences, supports "
    "decision-making, and learns from what actually happens."
)

doc.save(out)
print(out)
